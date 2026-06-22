import sys
from pathlib import Path

from docx import Document


def replace_paragraph_text(paragraph, text):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    if text:
        paragraph.add_run(text)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: update_section_4_7_10k_speedup.py <input.docx> <output.docx>")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(input_path))

    replacements = {
        155: "该实验场景下。结果表明，在当前QAT训练侧算子尚未充分融合的条件下，低比特rollout的推理收益需要在较长decode场景中才能完全覆盖训练侧额外开销。为进一步评估长序列生成场景下的端到端收益，本节将response length统一设为10k，并分别在两类假设下估算W8A8 rollout可带来的整网加速收益：",
        156: "（1）融合QAT算子优化60%假设下。若通过成熟的QAT融合算子减少中间张量读写与多次kernel launch，将训练侧QAT伪量化路径的额外开销压缩至当前实现的约60%，则模型端forward与backward的附加耗时将进一步下降。以当前profiling中L=4.3k时约3.4%的端到端收益为基准，并考虑decode阶段收益随输出长度近似线性累积，当response length提升至10k时，W8A8 rollout预计可获得约11.6%的端到端整网加速收益。该结果说明，在保留QAT训练路径的情况下，只要伪量化相关小算子得到充分融合，长序列decode累积收益即可覆盖训练侧额外开销，并转化为10%以上的整体吞吐提升。",
        157: "",
        158: "（2）去除QAT伪量化算子影响，仅保留W8A8 rollout路径假设。为估计低比特rollout推理自身的收益上限，我们进一步考虑训练阶段禁用QAT伪量化相关操作、仅保留W8A8低比特rollout推理路径的理想配置。在response length=4.3k、batch size=16的实测配置下，BF16基线单step训练时长为157.02s，仅保留W8A8 rollout后的单step训练时长降低至149.16s，对应理论端到端加速收益为5.27%。若将response length进一步提升至10k，并按decode收益随输出长度近似线性累积进行估算，则理论端到端加速收益可提升至约12.3%。",
        159: "",
        160: "上述两类假设表明，response length=10k时，W8A8 rollout已经具备超过10%的端到端整网加速潜力。其中，融合QAT算子优化60%假设对应约11.6%的可实现收益；去除QAT伪量化算子影响、仅保留低比特rollout路径的理想假设对应约12.3%的收益上限。二者的收益关系也符合系统开销逻辑：后者完全消除训练侧QAT伪量化开销，因此收益上限应高于仍保留部分QAT训练开销的融合算子优化场景。",
        161: "综上，W8A8 rollout已在推理侧体现出明确的prefill与decode加速收益，其中decode阶段收益会随生成长度线性累积；当前端到端收益受限的主要原因在于训练侧QAT伪量化路径仍存在额外计算开销。随着QAT伪量化与矩阵运算路径进一步融合优化，训练侧额外开销将进一步降低，低比特rollout的推理加速收益也更容易在10k及更长response length场景下转化为整体RL训练吞吐收益。",
        162: "",
    }

    for index, text in replacements.items():
        replace_paragraph_text(doc.paragraphs[index], text)

    doc.save(str(output_path))


if __name__ == "__main__":
    main()
