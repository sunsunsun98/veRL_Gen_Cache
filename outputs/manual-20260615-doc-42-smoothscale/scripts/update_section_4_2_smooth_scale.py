import copy
import sys
from pathlib import Path

from docx import Document
from docx.table import Table


def replace_paragraph_text(paragraph, text):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.add_run(text)


def set_table_text(table, rows):
    for r_idx, row_values in enumerate(rows):
        for c_idx, value in enumerate(row_values):
            table.cell(r_idx, c_idx).text = value


def move_after(anchor, *elements):
    cursor = anchor
    for element in elements:
        cursor.addnext(element)
        cursor = element


def renumber_tables(paragraph_text):
    replacements = [
        ("表4.2~4.3", "表4.3~4.4"),
        ("表4.2中的结果", "表4.3中的结果"),
        ("从表4.3中", "从表4.4中"),
        ("表4.2 模型openPangu-7B", "表4.3 模型openPangu-7B"),
        ("表4.3 模型Qwen3-30B", "表4.4 模型Qwen3-30B"),
        ("表4.4中QAT-W8", "表4.5中QAT-W8"),
        ("表4.5展示了", "表4.6展示了"),
        ("表4.4 模型openPangu-7B", "表4.5 模型openPangu-7B"),
        ("表4.5 模型Qwen3-30B", "表4.6 模型Qwen3-30B"),
        ("表4.6与表4.7", "表4.7与表4.8"),
        ("表4.6 模型openPangu-7B", "表4.7 模型openPangu-7B"),
        ("表4.7 模型Qwen3-30B", "表4.8 模型Qwen3-30B"),
        ("实验结果如表4.8所示", "实验结果如表4.9所示"),
        ("表4.8 模型openPangu-7B 数据集DAPO-Math-17K", "表4.9 模型openPangu-7B 数据集DAPO-Math-17K"),
        ("表4.8 不同规模模型下", "表4.10 不同规模模型下"),
        ("表4.8展示了不同模型规模下", "表4.10展示了不同模型规模下"),
        ("表4.9给出了", "表4.11给出了"),
        ("表4.9 Pangu-7B模型", "表4.11 Pangu-7B模型"),
        ("如表4.9所示", "如表4.11所示"),
        ("如表4.10所示", "如表4.12所示"),
        ("表4.10 Pangu-7B模型", "表4.12 Pangu-7B模型"),
        ("由表4.10可知", "由表4.12可知"),
    ]
    result = paragraph_text
    placeholders = []
    for index, (old, new) in enumerate(replacements):
        token = f"@@TABLE_RENUMBER_{index}@@"
        if old in result:
            result = result.replace(old, token)
            placeholders.append((token, new))
    for token, new in placeholders:
        result = result.replace(token, new)
    return result


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: update_section_4_2_smooth_scale.py <input.docx> <output.docx>")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(input_path))

    # 4.2 title and weight-scale analysis paragraph.
    replace_paragraph_text(doc.paragraphs[82], "可学习量化因子协同优化效果分析")
    replace_paragraph_text(
        doc.paragraphs[91],
        "此外，为优化权重量化网格划分并精准匹配动态量化区间，我们首先将weight scale（权重量化因子）设计为可学习参数，并通过消融实验验证其对QAT训练精度的增益。表4.1中train-qatw8-rollout-w8a8+ls（含可学习Scale）与train-qatw8-rollout-w8a8（不含可学习Scale）的对比结果表明，引入动态可学习的权重量化因子后，模型能够在训练过程中自适应调整量化步长，使权重分布与INT8量化网格更充分匹配，从而缓解权重量化带来的精度损失（Reward从0.444提升至0.451），并将训练与rollout之间的Pearson相关系数由0.9478提升至0.9723。该结果说明，weight scale不仅能够改善低比特权重量化的重构精度，也能够提升训练侧伪量化路径与推理侧真实量化路径的一致性。",
    )
    replace_paragraph_text(
        doc.paragraphs[92],
        "表4.1 模型openPangu-7B 数据集GSM8K下weight scale性能表现效果。（bf: bfloat; w: weight; a: activation; ls: learnable scale，下表同）表中结果表明，引入动态可学习的权重量化因子，能够缓解量化带来的精度损失，并改善训推一致性。",
    )

    # Renumber existing downstream table references before adding the new Table 4.2.
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = renumber_tables(text)
        if new_text != text:
            replace_paragraph_text(paragraph, new_text)

    # Add smooth-scale analysis after table 4.1, keeping the table format consistent with table 4.1.
    intro = doc.add_paragraph(
        "在此基础上，我们进一步评估smooth scale（量化平滑因子）对激活侧量化误差的抑制效果。与weight scale主要作用于权重量化步长不同，smooth scale面向W8A8 rollout中更突出的激活离群值问题，通过可学习的平滑因子对线性层输入激活与权重进行重参数化，将激活侧异常峰值平滑迁移至权重侧可吸收的尺度变化中，从而降低A8量化过程中的截断误差与舍入扰动。为验证该机制在复杂数学推理任务中的有效性，我们在DAPO-Math-17K数据集和openPangu-7B模型上进行了消融实验，结果如表4.2所示。"
    )
    intro.style = doc.paragraphs[91].style

    caption = doc.add_paragraph(
        "表4.2 模型openPangu-7B 数据集DAPO-Math-17K下smooth scale性能表现效果。（ss: smooth scale）表中结果表明，引入可学习量化平滑因子后，模型Reward与训推Pearson相关系数均得到提升，说明smooth scale能够有效缓解激活离群值导致的W8A8量化误差。"
    )
    caption.style = doc.paragraphs[92].style

    cloned_tbl = copy.deepcopy(doc.tables[0]._tbl)
    if len(cloned_tbl.tr_lst) > 5:
        cloned_tbl.remove(cloned_tbl.tr_lst[-1])
    smooth_table = Table(cloned_tbl, doc)
    set_table_text(
        smooth_table,
        [
            [" Methods", "Metrics", "Metrics"],
            [" Methods", "Reward（mean）", "Pearson_coef（train vs rollout）"],
            ["train-bf16-rollout-bf16", "0.309", "0.9986"],
            ["train-qatw8a8-rollout-w8a8", "0.252", "0.9802"],
            ["train-qatw8a8-rollout-w8a8+ss", "0.277", "0.9858"],
        ],
    )

    conclusion = doc.add_paragraph(
        "从实验结果可以看出，在未引入smooth scale时，train-qatw8a8-rollout-w8a8配置的Reward仅为0.252，Pearson相关系数为0.9802，说明激活侧A8量化误差会在复杂推理数据集上造成明显精度损失。加入smooth scale后，Reward提升至0.277，相较基础QAT W8A8配置提升0.025，Pearson相关系数同步提升至0.9858，表明可学习量化平滑因子能够有效降低激活离群值对rollout推理轨迹的扰动，并进一步增强训练侧伪量化与推理侧低比特执行之间的一致性。尽管该配置与BF16基线仍存在一定差距，但其在不改变W8A8整体量化形态的前提下显著改善了低比特rollout精度，为后续叠加随机舍入、敏感层回退等策略提供了更稳定的量化基础。"
    )
    conclusion.style = doc.paragraphs[91].style

    anchor_tbl = doc.tables[0]._tbl
    move_after(anchor_tbl, intro._p, caption._p, smooth_table._tbl, conclusion._p)

    doc.save(str(output_path))


if __name__ == "__main__":
    main()
